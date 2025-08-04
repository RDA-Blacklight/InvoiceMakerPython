from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

translations = {
    1:{
        "firstOptionList": "What would you like to? \n 1. Create New invoice \n",
        "innum": "What number invoice will this be? \n",
        "fabrication": "Fabrication/Installation \n Quantity(square foot):",
        "wasfab": "Was there Fabrication/Installation that occurred? \n 1.Yes \n 2.No \n",
        "template": "Template \n Quantity(square foot):",
        "wastemp": "Was there Templating that occured? \n 1.Yes \n 2.No \n",
        "wassink": "Was there any sinks? \n 1.Yes \n 2.No \n",
        "selsink": "Select one type of sink there was \n 1.Scoop Sink \n",
        "qtysink": "How many of these sinks was there? \n",
        "moresink": "Would you like to add another sink? \n 1.Yes \n 2.No \n",
        "sinkcut": "Was there any sink cutouts? \n 1.Yes \n 2.No \n",
        "selsinkcut": "What type of sink cutout was it? \n 1.Kitchen Sink Cut Out \n 2.Bath Sink Cut Out \n 3.Laundry Sink Cut Out \n",
        "qtysinkcut": "How many of these sink cuts where there? \n",
        "moresinkcut": "Would you like to add another sink cut? \n 1.Yes \n 2.No \n",
        "rate": "Is this rate correct?",
        "norate": "What is the rate? \n $",
        "yesno": " 1.Yes \n 2.No \n",
        "address": "What is the address of the location \n",

        "custom": "Would you like to add something custom? \n 1.Yes \n 2.No \n",
        "customqty": "How many of this custom item is there? \n",
        "iscustomunit": "Would you like to put units?",
        "selcustomunit": "1. Sq.Ft \n 2. Each",
        "iscustprice": "How much does this cost? \n $",
        "ismessage": "Would you like to add a custom message? \n 1.Yes \n 2.No \n",
        "message": "What would you like to say? \n",
        "customname": "What would you like to call this custom item? \n",
        "morecustom": "Would you like to add another custom item? \n 1.Yes \n 2.No \n"
    },
    2:{
        "firstOptionList": "O que você gostaria de fazer? \n 1. Criar Novo fatura(invoice) \n",
        "innum": "Qual será o número da fatura? \n",
        "fabrication": "Fabricação/Instalação \n Quantidade(pé quadrado):",
        "wasfab": "Houve Fabricação/Instalação que aconteceu? \n 1.Sim \n 2.Não \n",
        "template": "Modelo(Template) \n Quantidade(pé quadrado):",
        "wastemp": "Houve alguma modelagem(Template) que ocorreu? \n 1.Sim \n 2.Não \n",
        "wassink": "Tinha alguma pia? \n 1.Sim \n 2.Não \n",
        "selsink": "Selecione um tipo de pia que havia lá \n 1.Scoop Sink \n",
        "qtysink": "Quantas dessas pias havia \n",
        "moresink": "Você gostaria de adicionar outra pia? \n 1.Sim \n 2.Não \n",
        "sinkcut": "Houve algum corte para a pia(Sink cutouts)? \n 1.Sim \n 2.Não \n",
        "selsinkcut": "Que tipo de corte de pia era? \n 1. Recorte de Pia de Cozinha \n 2. Recorte de Pia de Banho \n 3. Recorte de Pia de Lavanderia \n",
        "qtysinkcut": "Quantos cortes de pia haviam aqui? \n",
        "moresinkcut": "Você gostaria de adicionar mais um corte para a pia? \n 1.Sim \n 2.Não \n",
        "rate": "Esta taxa(rate) está correta?",
        "norate": "Qual é a taxa? \n $",
        "yesno": " 1.Sim \n 2.Náo \n",
        "address": "Qual é o endereço do local? \n",

        "custom": "Você gostaria de adicionar algo personalizado? \n 1.Sim \n 2.Náo \n",
        "customqty": "Quantos desses itens personalizados existem? \n",
        "iscustomunit": "Você gostaria de colocar unidades?",
        "selcustomunit": "1. pés quadrados \n 2. Cada",
        "iscustprice": "Quanto custa isso? \n $",
        "ismessage": "Você gostaria de adicionar uma mensagem personalizada?? \n 1.Sim \n 2.Náo \n",
        "message": "O que você gostaria de dizer? \n",
        "customname": "Como você gostaria de chamar este item personalizado? \n",
        "morecustom": "Você gostaria de adicionar outro item personalizado? \n 1.Sim \n 2.Náo \n"


    }
}

lang = int(input("Select a language \n Selecione um idioma \n 1. English 2. Português \n"))


def invoice(info, placeholder, col1holder, col2holder, col3holder, custommessage):

    doc = Document("INVOICE NUMBER.docx")

    table = doc.tables[1]
    table.style = 'Table Grid'

    hdr_cell = table.rows[0].cells
    hdr_cell[0].text = 'Item'
    hdr_cell[1].text = 'Quantity'
    hdr_cell[2].text = 'Rate'
    hdr_cell[3].text = 'Amount'

    placeholder.append("{{total}}")
    total = 0

    for n in range(len(col1holder)):
        new_row = table.add_row()
        new_row.cells[0].text = str(col1holder[n])
        new_row.cells[1].text = str(col2holder[n])
        new_row.cells[2].text = ("$" + str(col3holder[n]))
        new_row.cells[3].text = ("$" + str(col2holder[n] * col3holder[n]))
        if str(col1holder[n]) in ("Fabrication/Installation", "Template"):
            new_row.cells[1].text += " Sq.Ft"
        if str(col1holder[n]) in ("Fabrication/Installation", "Template"):
            new_row.cells[2].text += " per Sq.Ft"
        elif "Sink" in str(col1holder[n]):
            new_row.cells[2].text += " each"

        total += (col2holder[n] * col3holder[n])

    info.append(total)

    table = doc.tables[1]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                for a in range(len(placeholder)):
                    if placeholder[a] in cell_text:
                        new_text = cell_text.replace(placeholder[a], str(info[a]))
                        cell.text = new_text

    if custommessage and custommessage[0] == 1:
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(custommessage[1])

    doc.save("Invoice" + str(info[0]) + ".docx")

    docx_path = "Invoice" + str(info[0]) + ".docx"
    pdf_path = "Invoice" + str(info[0]) + ".pdf"
    convert(docx_path, pdf_path)



def template():
    print("\033[2J\033[H")

    info = []
    placeholder = []

    col1holder = []
    col2holder = []
    col3holder = []

    custommessage = []

    invoiceNumber = int(input(translations[lang]["innum"]))
    info.append(invoiceNumber)
    placeholder.append("{{Number}}")
    print("\033[2J\033[H")

    address = input(translations[lang]["address"])
    info.append(address)
    placeholder.append("{{address}}")
    print("\033[2J\033[H")

    wasfab = int(input(translations[lang]["wasfab"]))
    print("\033[2J\033[H")
    if wasfab == 1:
        fabqty = int(input(translations[lang]["fabrication"]))
        print("\033[2J\033[H")
        col2holder.append(fabqty)
        col1holder.append("Fabrication/Installation")
        rate = 18
        print("\033[2J\033[H")
        israte = int(input(translations[lang]["rate"] + " $" + str(rate) + "\n" + (translations[lang]["yesno"])))
        if israte == 2:
            print("\033[2J\033[H")
            rate = int(input(translations[lang]["norate"]))
            print("\033[2J\033[H")
        col3holder.append(rate)

    print("\033[2J\033[H")
    wastemp = int(input(translations[lang]["wastemp"]))
    print("\033[2J\033[H")
    if wastemp == 1:
        tempqty = int(input(translations[lang]["template"]))
        print("\033[2J\033[H")
        col2holder.append(tempqty)
        col1holder.append("Template")
        rate = 2
        israte = int(input(translations[lang]["rate"] + " $" + str(rate) + "\n" + (translations[lang]["yesno"])))
        print("\033[2J\033[H")
        if israte == 2:
            rate = int(input(translations[lang]["norate"]))
            print("\033[2J\033[H")
        col3holder.append(rate)

    wassink = int(input(translations[lang]["wassink"]))
    print("\033[2J\033[H")
    while wassink == 1:
        selsink = int(input(translations[lang]["selsink"]))
        print("\033[2J\033[H")
        if selsink == 1:
            scoopsinkqty = int(input(translations[lang]["qtysink"]))
            print("\033[2J\033[H")
            col2holder.append(scoopsinkqty)
            col1holder.append("Scoop Sink")
            rate = 35
            israte = int(input(translations[lang]["rate"] + " $" + str(rate) + "\n" + (translations[lang]["yesno"])))
            print("\033[2J\033[H")
            if israte == 2:
                rate = int(input(translations[lang]["norate"]))
                print("\033[2J\033[H")
            col3holder.append(rate)
        wassink = int(input(translations[lang]["moresink"]))
        print("\033[2J\033[H")


    sinkcut = int(input(translations[lang]["sinkcut"]))
    print("\033[2J\033[H")
    while sinkcut == 1:
        selsinkcut = int(input(translations[lang]["selsinkcut"]))
        print("\033[2J\033[H")
        if selsinkcut == 1:
            kitchsinkcut = int(input(translations[lang]["qtysinkcut"]))
            print("\033[2J\033[H")
            col2holder.append(kitchsinkcut)
            col1holder.append("Kitchen Sink Cut Out")

            rate = 75
            israte = int(input(translations[lang]["rate"] + " $" + str(rate) + "\n" + (translations[lang]["yesno"])))
            print("\033[2J\033[H")
            if israte == 2:
                rate = int(input(translations[lang]["norate"]))
                print("\033[2J\033[H")
            col3holder.append(rate)

        if selsinkcut == 2:
            bathsinkcut = int(input(translations[lang]["qtysinkcut"]))
            print("\033[2J\033[H")
            col2holder.append(bathsinkcut)
            col1holder.append("Bath Sink Cut Out")

            rate = 50
            israte = int(input(translations[lang]["rate"] + " $" + str(rate) + "\n" + (translations[lang]["yesno"])))
            print("\033[2J\033[H")
            if israte == 2:
                rate = int(input(translations[lang]["norate"]))
                print("\033[2J\033[H")
            col3holder.append(rate)

        if selsinkcut == 3:
            laundrysinkcut = int(input(translations[lang]["qtysinkcut"]))
            print("\033[2J\033[H")
            col2holder.append(laundrysinkcut)
            col1holder.append("Laundry Sink Cut Out")

            rate = 75
            israte = int(input(translations[lang]["rate"] + " $" + str(rate) + "\n" + (translations[lang]["yesno"])))
            print("\033[2J\033[H")
            if israte == 2:
                rate = int(input(translations[lang]["norate"]))
                print("\033[2J\033[H")
            col3holder.append(rate)
        sinkcut = int(input(translations[lang]["moresinkcut"]))
        print("\033[2J\033[H")


    wantcustom = int(input(translations[lang]["custom"]))
    print("\033[2J\033[H")
    while wantcustom == 1:
        name = input(translations[lang]["customname"])
        print("\033[2J\033[H")
        col1holder.append(name)
        customqty = int(input(translations[lang]["customqty"]))
        print("\033[2J\033[H")
        col2holder.append(customqty)
        customprice = int(input(translations[lang]["iscustprice"]))
        print("\033[2J\033[H")
        col3holder.append(customprice)
        wantcustom = int(input(translations[lang]["morecustom"]))
        print("\033[2J\033[H")

    ismessage = int(input(translations[lang]["ismessage"]))
    print("\033[2J\033[H")
    if ismessage == 1:
        custommessage.append(1)
        message = input(translations[lang]["message"])
        print("\033[2J\033[H")
        custommessage.append(message)

    invoice(info, placeholder, col1holder, col2holder, col3holder, custommessage)
def menu():
    print("\033[2J\033[H")
    choice = int(input(translations[lang]["firstOptionList"]))
    if choice == 1:
        template()

menu()